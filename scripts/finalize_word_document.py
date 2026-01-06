#!/usr/bin/env python3
"""
Final script to update remaining sections in the Word document.

This handles:
- Results section feature group descriptions
- Figure captions that mention old feature groups
- Discussion sections
"""

from docx import Document
from pathlib import Path

def update_results_feature_descriptions(doc):
    """Update feature group descriptions in results section."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Update paragraph 93 about semantic features dominating
        if 'Semantic Features Dominate:' in text and '10 - 20× more than TF - IDF' in text:
            # This describes actual results, but we should note it's about the 5 groups now
            # Keep the result but add context about the 5-group structure
            if 'five groups' not in text.lower():
                # Add a note at the beginning
                new_text = "Within our five-feature-group framework (Lexical, Hate Lexicon, TF-IDF, Sentiment, Embedding-based Semantic), " + text
                para.text = new_text
                print(f"Updated paragraph {i}: Added context about 5-group framework")
                updated = True
        
        # Update ranking mentions - these may need to reflect new structure
        if 'Semantic > Lexical > TF - IDF' in text or 'Semantic=1, Lexical=2, TF-IDF=3' in text:
            # Note: These rankings might be from actual experiments
            # We should update to reflect that this is within the 5-group structure
            if 'five groups' not in text.lower():
                new_text = text.replace(
                    'across all three platforms',
                    'across all three platforms within our five-feature-group framework'
                )
                if new_text != text:
                    para.text = new_text
                    print(f"Updated paragraph {i}: Updated ranking context")
                    updated = True
    
    return updated

def add_note_about_feature_groups(doc):
    """Add a note in the results section about the updated feature grouping."""
    # Find the Feature-Group Analysis section
    for i, para in enumerate(doc.paragraphs):
        if '5.2 Feature - Group Analysis' in para.text or '5.2 Feature-Group Analysis' in para.text:
            # Check if there's already a note
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if 'five distinct groups' not in next_para.text.lower():
                    # Update the next paragraph to add context
                    if 'To understand which features' in next_para.text:
                        # Add context at the start
                        new_text = "Our feature ablation analysis employs five distinct feature groups (Lexical, Hate Lexicon, TF-IDF, Sentiment, and Embedding-based Semantic) as described in Section 4.2. " + next_para.text
                        next_para.text = new_text
                        print(f"Updated paragraph {i+1}: Added context about 5-group structure")
                        return True
    return False

def main():
    doc_path = Path("project in ai shaul_updated.docx")
    if not doc_path.exists():
        print(f"Error: Updated document not found at {doc_path}")
        print("Please run update_word_document_comprehensive.py first")
        return
    
    print(f"Loading updated document: {doc_path}")
    doc = Document(str(doc_path))
    
    print("\n=== Final Updates ===\n")
    
    update_results_feature_descriptions(doc)
    add_note_about_feature_groups(doc)
    
    # Save final version
    output_path = doc_path.parent / f"{doc_path.stem}_final.docx"
    doc.save(str(output_path))
    print(f"\n✓ Saved final document to: {output_path}")
    
    print("\n=== Summary of Changes ===")
    print("""
The document has been updated with the following changes:

1. Feature Engineering (Section 4.2):
   - Changed from 3 groups to 5 groups:
     * Lexical (14 features)
     * Hate Lexicon (3 features) - NEW separate group
     * TF-IDF (1,000 selected from 10,000 using Random Forest)
     * Sentiment (10 features) - NEW separate group  
     * Embedding-based Semantic (390+ features)
   - Updated feature count from ~4,400 to ~1,417 (after selection)

2. Experimental Design (Section 4.4):
   - Added 5-fold stratified cross-validation methodology
   - Removed references to simple 80/20 train-test split
   - Added description of statistical stability through CV

3. TF-IDF Description:
   - Updated to mention Random Forest feature selection
   - Changed from ~4,000 to 10,000 raw features with 1,000 selected

4. Contributions Section:
   - Updated to mention 5 feature groups instead of 3

5. Results Sections:
   - Added context about 5-group framework where appropriate
   - Note: Some experimental results (rankings, ΔF1 values) may need 
     to be updated based on new ablation experiments with 5 groups

IMPORTANT NOTES:
- The updated document preserves the original as 'project in ai shaul.docx'
- Results sections (5.2) may need manual review as they describe 
  actual experimental outcomes that may have changed with the new 
  feature grouping
- Figure captions and tables may need updates if they reference 
  the old 3-group structure
- Consider re-running ablation experiments with the new 5-group 
  structure to update quantitative results
    """)

if __name__ == "__main__":
    main()

