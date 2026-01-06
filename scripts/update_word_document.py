#!/usr/bin/env python3
"""
Script to update the Word document with new feature grouping structure and methodology.

Updates:
1. Feature Engineering section (4.2) - Change from 3 groups to 5 groups
2. Experimental Design section (4.4) - Add 5-fold CV methodology
3. Update TF-IDF description to mention Random Forest selection
"""

from docx import Document
from pathlib import Path

def update_feature_engineering_section(doc):
    """Update section 4.2 Feature Engineering with new 5-group structure."""
    
    # Find the paragraph with "We organize features into three conceptual groups"
    for i, para in enumerate(doc.paragraphs):
        if 'organize features into three conceptual groups' in para.text.lower():
            # Replace the entire feature engineering section
            new_text = """We organize features into five distinct groups to support comprehensive ablation analysis and understand feature contributions:

1. Lexical Features: Surface-level text properties capturing explicit patterns and stylistic characteristics. These include word count, average word length, uppercase letter ratio (indicating shouting or emphasis), exclamation mark ratio, question mark ratio, period ratio, punctuation count, multiple punctuation patterns, character count, repeated character patterns, unique word count, word repetition, longest and shortest word lengths, and language complexity index (Flesch-Kincaid readability score). Lexical features help detect explicit slurs, derogatory phrases, intensified punctuation patterns, and stylistic markers of toxic content. This group contains 14 features.

2. Hate Lexicon Features: Direct matching against a curated hate speech lexicon. These features include hate word count (absolute frequency of lexicon matches), hate word ratio (proportion of words that match the lexicon), and a binary flag indicating presence of any hate word. The lexicon is sourced from external datasets and provides explicit detection of known offensive terms. This group contains 3 features.

3. TF-IDF Features: Term Frequency-Inverse Document Frequency vectors capturing word importance relative to the corpus. We extract unigrams and bigrams from up to 10,000 potential features, then apply Random Forest-based feature selection to identify the top 1,000 most discriminative features. This supervised selection process ensures that only the most informative n-grams are retained, significantly improving model performance and reducing dimensionality. The TF-IDF transformation uses sublinear term frequency scaling: tfidf(w,d) = (1 + log(tf(w,d))) · log(N/df(w)). This group contains 1,000 selected features.

4. Sentiment Features: Emotional tone and subjectivity analysis using TextBlob. These include sentiment polarity (ranging from -1 to +1), sentiment subjectivity (0 to 1), sentiment magnitude (absolute polarity), sentiment category flags (negative, neutral, positive), extreme sentiment indicators (very negative, very positive), and subjectivity category flags (highly subjective, highly objective). Sentiment features capture the emotional valence and intensity of text, which are strong indicators of hate speech patterns. This group contains 10 features.

5. Embedding-based Semantic Features: Deep semantic understanding through sentence embeddings and linguistic analysis. These include 384-dimensional sentence embeddings from all-MiniLM-L6-v2, part-of-speech ratios (noun, verb, adjective), pronoun usage patterns (they, us, you ratios), hate lexicon similarity (cosine similarity to mean hate lexicon embedding), embedding distance to neutral content, and language complexity metrics. Semantic features help identify stereotypes, threats, and hostility without explicit slurs, capturing outgroup generalization and contextual understanding essential for detecting coded language and sophisticated hate speech. This group contains approximately 390+ features.

The final feature space concatenates all groups: X = [X_lexical, X_hate_lexicon, X_tfidf, X_sentiment, X_semantic], providing approximately 1,417 features per sample (down from ~4,400 raw features after TF-IDF selection). This mixed representation allows models to learn both platform-specific slang and transferable toxicity patterns across domains."""
            
            para.text = new_text
            print(f"Updated paragraph {i}: Feature Engineering section")
            return i
    
    print("Warning: Could not find feature engineering section to update")
    return None

def update_experimental_design_section(doc):
    """Update section 4.4 Experimental Design to mention 5-fold CV."""
    
    # Find the paragraph with "Cross - Dataset Evaluation"
    for i, para in enumerate(doc.paragraphs):
        if 'Cross - Dataset Evaluation:' in para.text or 'Cross-Dataset Evaluation:' in para.text:
            # Check if 5-fold CV is already mentioned
            if '5-fold' in para.text.lower() or 'five-fold' in para.text.lower():
                print(f"Paragraph {i} already mentions 5-fold CV")
                continue
            
            # Update the paragraph to include 5-fold CV methodology
            original_text = para.text
            new_text = """Cross-Dataset Evaluation: For statistical stability and robustness, we employ 5-fold stratified cross-validation for each dataset. Each dataset is divided into 5 folds, ensuring that each fold maintains the original class distribution. For each platform, we train a separate logistic regression classifier and evaluate it using all 5 folds of each dataset (including its own held-out folds). This yields comprehensive train→test combinations with aggregated metrics (mean ± standard deviation) across folds, providing statistically stable performance estimates. We record Accuracy, Precision, Recall, F1-score, ROC-AUC, and Average Precision for each combination, with results aggregated across all 5 folds."""
            
            para.text = new_text
            print(f"Updated paragraph {i}: Experimental Design section")
            
            # Also check the next few paragraphs for related content
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                if 'train a separate' in next_para.text.lower() and '80%' in next_para.text:
                    # Update to mention 5-fold CV instead of simple 80/20 split
                    next_para.text = "This 5-fold cross-validation approach ensures that performance metrics reflect true generalization rather than single-split variability, providing more reliable estimates of cross-platform transferability."
                    print(f"Updated paragraph {j}: Removed outdated 80/20 split reference")
            return i
    
    print("Warning: Could not find experimental design section to update")
    return None

def update_feature_ablation_section(doc):
    """Update feature ablation experiments to reflect 5 groups."""
    
    for i, para in enumerate(doc.paragraphs):
        if 'Feature - Group Ablation:' in para.text or 'Feature-Group Ablation:' in para.text:
            # Update to mention 5 groups instead of 3
            original_text = para.text
            if 'five model variants' in original_text.lower() or '5 model variants' in original_text.lower():
                # Already updated or correct
                continue
            
            new_text = """Feature-Group Ablation: To quantify feature importance, we train multiple model variants per dataset using different feature group combinations: (A) FULL - all five feature groups, (B) LEXICAL_ONLY - only lexical features, (C) HATE_LEXICON_ONLY - only hate lexicon features, (D) TFIDF_ONLY - only TF-IDF features, (E) SENTIMENT_ONLY - only sentiment features, (F) SEMANTIC_ONLY - only embedding-based semantic features, and leave-one-group-out ablations (G-K) removing each group individually. We measure feature group contribution using ΔF1 = F1(FULL) - F1(without group), where higher ΔF1 indicates greater importance."""
            
            para.text = new_text
            print(f"Updated paragraph {i}: Feature ablation section")
            return i
    
    return None

def main():
    doc_path = Path("project in ai shaul.docx")
    if not doc_path.exists():
        print(f"Error: Document not found at {doc_path}")
        return
    
    print(f"Loading document: {doc_path}")
    doc = Document(str(doc_path))
    
    print("\nUpdating Feature Engineering section (4.2)...")
    update_feature_engineering_section(doc)
    
    print("\nUpdating Experimental Design section (4.4)...")
    update_experimental_design_section(doc)
    
    print("\nUpdating Feature Ablation section...")
    update_feature_ablation_section(doc)
    
    # Save the updated document
    output_path = doc_path.parent / f"{doc_path.stem}_updated.docx"
    doc.save(str(output_path))
    print(f"\n✓ Saved updated document to: {output_path}")
    print(f"Original document preserved at: {doc_path}")

if __name__ == "__main__":
    main()

