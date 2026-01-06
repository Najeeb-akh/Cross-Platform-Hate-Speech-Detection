#!/usr/bin/env python3
"""
Comprehensive script to update the Word document with new feature grouping structure.

This script updates:
1. Feature Engineering section (4.2) - Change from 3 groups to 5 groups
2. Experimental Design section (4.4) - Add 5-fold CV methodology  
3. Results sections (5.2) - Update feature ablation descriptions
4. Abstract and Contributions - Update feature group mentions
5. All references to old feature counts (4400, 4000) to new counts (1417, 1000 selected)
"""

from docx import Document
from pathlib import Path
import re

def update_paragraph_text(para, old_pattern, new_text, description=""):
    """Update paragraph text if pattern matches."""
    if old_pattern.lower() in para.text.lower():
        para.text = new_text
        if description:
            print(f"  ✓ {description}")
        return True
    return False

def update_feature_engineering_section(doc):
    """Update section 4.2 Feature Engineering with new 5-group structure."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        if 'organize features into three conceptual groups' in para.text.lower():
            new_text = """We organize features into five distinct groups to support comprehensive ablation analysis and understand feature contributions:

1. Lexical Features: Surface-level text properties capturing explicit patterns and stylistic characteristics. These include word count, average word length, uppercase letter ratio (indicating shouting or emphasis), exclamation mark ratio, question mark ratio, period ratio, punctuation count, multiple punctuation patterns, character count, repeated character patterns, unique word count, word repetition, longest and shortest word lengths, and language complexity index (Flesch-Kincaid readability score). Lexical features help detect explicit slurs, derogatory phrases, intensified punctuation patterns, and stylistic markers of toxic content. This group contains 14 features.

2. Hate Lexicon Features: Direct matching against a curated hate speech lexicon. These features include hate word count (absolute frequency of lexicon matches), hate word ratio (proportion of words that match the lexicon), and a binary flag indicating presence of any hate word. The lexicon is sourced from external datasets and provides explicit detection of known offensive terms. This group contains 3 features.

3. TF-IDF Features: Term Frequency-Inverse Document Frequency vectors capturing word importance relative to the corpus. We extract unigrams and bigrams from up to 10,000 potential features, then apply Random Forest-based feature selection to identify the top 1,000 most discriminative features. This supervised selection process ensures that only the most informative n-grams are retained, significantly improving model performance and reducing dimensionality. The TF-IDF transformation uses sublinear term frequency scaling: tfidf(w,d) = (1 + log(tf(w,d))) · log(N/df(w)). This group contains 1,000 selected features.

4. Sentiment Features: Emotional tone and subjectivity analysis using TextBlob. These include sentiment polarity (ranging from -1 to +1), sentiment subjectivity (0 to 1), sentiment magnitude (absolute polarity), sentiment category flags (negative, neutral, positive), extreme sentiment indicators (very negative, very positive), and subjectivity category flags (highly subjective, highly objective). Sentiment features capture the emotional valence and intensity of text, which are strong indicators of hate speech patterns. This group contains 10 features.

5. Embedding-based Semantic Features: Deep semantic understanding through sentence embeddings and linguistic analysis. These include 384-dimensional sentence embeddings from all-MiniLM-L6-v2, part-of-speech ratios (noun, verb, adjective), pronoun usage patterns (they, us, you ratios), hate lexicon similarity (cosine similarity to mean hate lexicon embedding), embedding distance to neutral content, and language complexity metrics. Semantic features help identify stereotypes, threats, and hostility without explicit slurs, capturing outgroup generalization and contextual understanding essential for detecting coded language and sophisticated hate speech. This group contains approximately 390+ features.

The final feature space concatenates all groups: X = [X_lexical, X_hate_lexicon, X_tfidf, X_sentiment, X_semantic], providing approximately 1,417 features per sample (down from ~4,400 raw features after TF-IDF selection). This mixed representation allows models to learn both platform-specific slang and transferable toxicity patterns across domains."""
            
            para.text = new_text
            print(f"Updated paragraph {i}: Feature Engineering section (4.2)")
            updated = True
            break
    
    return updated

def update_experimental_design_section(doc):
    """Update section 4.4 Experimental Design to mention 5-fold CV."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if 'Cross - Dataset Evaluation:' in text or 'Cross-Dataset Evaluation:' in text:
            # Check if already updated
            if '5-fold' in text.lower() or 'five-fold' in text.lower():
                continue
            
            new_text = """Cross-Dataset Evaluation: For statistical stability and robustness, we employ 5-fold stratified cross-validation for each dataset. Each dataset is divided into 5 folds, ensuring that each fold maintains the original class distribution. For each platform, we train a separate logistic regression classifier and evaluate it using all 5 folds of each dataset (including its own held-out folds). This yields comprehensive train→test combinations with aggregated metrics (mean ± standard deviation) across folds, providing statistically stable performance estimates. We record Accuracy, Precision, Recall, F1-score, ROC-AUC, and Average Precision for each combination, with results aggregated across all 5 folds."""
            
            para.text = new_text
            print(f"Updated paragraph {i}: Experimental Design - Cross-Dataset Evaluation")
            updated = True
            
            # Update the next paragraph if it mentions 80/20 split
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '80%' in next_para.text and '20%' in next_para.text:
                    next_para.text = "This 5-fold cross-validation approach ensures that performance metrics reflect true generalization rather than single-split variability, providing more reliable estimates of cross-platform transferability."
                    print(f"Updated paragraph {i+1}: Removed outdated 80/20 split reference")
            break
    
    return updated

def update_feature_ablation_section(doc):
    """Update feature ablation descriptions to reflect 5 groups."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # Update mentions of "five model variants" or ablation experiments
        if 'Feature - Group Ablation:' in text or 'Feature-Group Ablation:' in text:
            if 'five model variants' not in text.lower() and '5 model variants' not in text.lower():
                new_text = """Feature-Group Ablation: To quantify feature importance, we train multiple model variants per dataset using different feature group combinations: (A) FULL - all five feature groups, (B) LEXICAL_ONLY - only lexical features, (C) HATE_LEXICON_ONLY - only hate lexicon features, (D) TFIDF_ONLY - only TF-IDF features, (E) SENTIMENT_ONLY - only sentiment features, (F) SEMANTIC_ONLY - only embedding-based semantic features, and leave-one-group-out ablations (G-K) removing each group individually. We measure feature group contribution using ΔF1 = F1(FULL) - F1(without group), where higher ΔF1 indicates greater importance."""
                para.text = new_text
                print(f"Updated paragraph {i}: Feature-Group Ablation description")
                updated = True
    
    return updated

def update_contributions_section(doc):
    """Update contributions section to mention 5 feature groups."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if 'Comprehensive feature - group ablation analysis' in text or 'Comprehensive feature-group ablation analysis' in text:
            if 'five' not in text.lower() and '5' not in text:
                # Update to mention 5 groups
                new_text = text.replace(
                    'quantifying contributions of lexical, semantic, and TF - IDF features',
                    'quantifying contributions of lexical, hate lexicon, TF-IDF, sentiment, and embedding-based semantic features across five distinct feature groups'
                )
                para.text = new_text
                print(f"Updated paragraph {i}: Contributions section")
                updated = True
                break
    
    return updated

def update_feature_count_references(doc):
    """Update references to old feature counts (4400, 4000) to new counts."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # Replace ~4000 or 4000 with 10,000 (raw) and mention 1,000 selected
        if '~4000' in text or '4000 dimensions' in text or 'approximately 4400' in text:
            new_text = text.replace('~4000', '10,000 (with 1,000 selected)')
            new_text = new_text.replace('4000 dimensions', '10,000 dimensions (1,000 selected)')
            new_text = new_text.replace('approximately 4400', 'approximately 1,417')
            new_text = new_text.replace('~4400', '~1,417')
            if new_text != text:
                para.text = new_text
                print(f"Updated paragraph {i}: Feature count references")
                updated = True
    
    return updated

def update_results_sections(doc):
    """Update results sections that mention feature groups."""
    updated = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # Update mentions of "three groups" or specific group comparisons
        if 'lexical, semantic, and TF - IDF' in text or 'lexical, semantic, and TF-IDF' in text:
            new_text = text.replace(
                'lexical, semantic, and TF - IDF',
                'lexical, hate lexicon, TF-IDF, sentiment, and embedding-based semantic'
            )
            new_text = new_text.replace(
                'lexical, semantic, and TF-IDF',
                'lexical, hate lexicon, TF-IDF, sentiment, and embedding-based semantic'
            )
            if new_text != text:
                para.text = new_text
                print(f"Updated paragraph {i}: Results section feature group mention")
                updated = True
    
    return updated

def main():
    doc_path = Path("project in ai shaul.docx")
    if not doc_path.exists():
        print(f"Error: Document not found at {doc_path}")
        return
    
    print(f"Loading document: {doc_path}")
    doc = Document(str(doc_path))
    
    print("\n=== Updating Document Sections ===\n")
    
    updates = []
    updates.append(("Feature Engineering (4.2)", update_feature_engineering_section(doc)))
    updates.append(("Experimental Design (4.4)", update_experimental_design_section(doc)))
    updates.append(("Feature Ablation", update_feature_ablation_section(doc)))
    updates.append(("Contributions", update_contributions_section(doc)))
    updates.append(("Feature Count References", update_feature_count_references(doc)))
    updates.append(("Results Sections", update_results_sections(doc)))
    
    print("\n=== Summary ===")
    for name, was_updated in updates:
        status = "✓ Updated" if was_updated else "○ No changes needed"
        print(f"{status}: {name}")
    
    # Save the updated document
    output_path = doc_path.parent / f"{doc_path.stem}_updated.docx"
    doc.save(str(output_path))
    print(f"\n✓ Saved updated document to: {output_path}")
    print(f"Original document preserved at: {doc_path}")

if __name__ == "__main__":
    main()

